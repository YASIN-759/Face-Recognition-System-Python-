
import pymysql
from tkinter import *
from tkinter import ttk, messagebox
from docx import Document
from docx.shared import Inches
import os

class PresentStudentsFrame:
    def __init__(self, root):
        self.root = root
        self.root.title("Present Students")
        self.root.attributes('-fullscreen', True)
        self.root.bind("<Escape>", lambda e: self.exit_fullscreen())
        self.root.configure(bg="#0A192F")

        # Clear existing widgets
        for widget in self.root.winfo_children():
            widget.destroy()

        # Heading
        Label(self.root, text="Present Students", font=("Arial", 24, "bold"),
              fg="#00FFFF", bg="#0A192F").pack(pady=20)

        # Treeview table
        self.table_frame = Frame(self.root, bg="#0A192F")
        self.table_frame.pack(pady=10, fill=BOTH, expand=True)

        self.columns = ("Roll No", "Name", "Department", "Timestamp")
        self.tree = ttk.Treeview(self.table_frame, columns=self.columns, show="headings")
        for col in self.columns:
            self.tree.heading(col, text=col)
            self.tree.column(col, anchor="center", width=150)
        self.tree.pack(fill=BOTH, expand=True)

        # Buttons
        btn_frame = Frame(self.root, bg="#0A192F")
        btn_frame.pack(pady=10)

        Button(btn_frame, text="Export to Word", font=("Arial", 14, "bold"), bg="#00FFFF", fg="#000000",
               command=self.export_to_word).grid(row=0, column=0, padx=20)

        Button(btn_frame, text="Clear Attendance History", font=("Arial", 14, "bold"), bg="#FF4C4C", fg="white",
               command=self.clear_attendance).grid(row=0, column=1, padx=20)

        Button(btn_frame, text="Go to Main", font=("Arial", 14, "bold"), bg="red", fg="white",
               command=self.go_to_main).grid(row=0, column=2, padx=20)

        # Load data on start
        self.load_present_students()

    def exit_fullscreen(self):
        self.root.attributes('-fullscreen', False)
        self.root.geometry("1351x800")

    def load_present_students(self):
        """Load and display present students from the 'attendance' table in MySQL."""
        try:
            conn = pymysql.connect(host="localhost", user="root", password="",
                                   database="fac", port=3306)
            cursor = conn.cursor()
            cursor.execute("SELECT roll, name, department, timestamp FROM attendance WHERE status='present'")
            self.records = cursor.fetchall()
            conn.close()

            self.tree.delete(*self.tree.get_children())
            for row in self.records:
                self.tree.insert("", END, values=row)

        except Exception as e:
            print("Error fetching attendance:", e)

    def export_to_word(self):
        """Export present students to a Word document."""
        try:
            if not self.records:
                print("No records to export.")
                return

            doc = Document()
            doc.add_heading("Present Students Report", 0)

            table = doc.add_table(rows=1, cols=len(self.columns))
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(self.columns):
                hdr_cells[i].text = col

            for row in self.records:
                row_cells = table.add_row().cells
                for i, item in enumerate(row):
                    row_cells[i].text = str(item)

            doc.save("Present_Students_Report.docx")
            print("Report saved as 'Present_Students_Report.docx'.")

        except Exception as e:
            print("Error exporting to Word:", e)

    def clear_attendance(self):
        """Delete all records from the attendance table."""
        confirm = messagebox.askyesno("Confirm Delete", "Are you sure you want to delete all attendance records?", parent=self.root)
        if confirm:
            try:
                conn = pymysql.connect(host="localhost", user="root", password="",
                                       database="fac", port=3306)
                cursor = conn.cursor()
                cursor.execute("DELETE FROM attendance")
                conn.commit()
                conn.close()

                self.load_present_students()
                messagebox.showinfo("Deleted", "All attendance records have been cleared.", parent=self.root)
            except Exception as e:
                messagebox.showerror("Error", f"Failed to clear attendance records: {str(e)}", parent=self.root)

    def go_to_main(self):
        """Return to the main UI by launching main_ui.py."""
        self.root.destroy()
        os.system("python main_ui.py")

# Main entry point
def main(root):
    PresentStudentsFrame(root)
